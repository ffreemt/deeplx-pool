"""Generate docx for source_text, trtext, alt_text."""
# pylint: disable=too-many-locals, broad-exception-caught, too-many-branches, too-many-statements,

import asyncio
import os
from itertools import zip_longest
from pathlib import Path
from secrets import token_hex
from time import monotonic
from typing import List, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX

from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from loadtext import loadtext
from loguru import logger
from rich.console import Console
from ycecream import y

from deeplx_pool.batch_deeplx_tr import batch_deeplx_tr

# from deeplx_pool.batch_newapi_tr import batch_newapi_tr
from deeplx_pool.duration_human import duration_human

console = Console()
y.configure(
    sln=1,
    st=1,
    e=os.getenv("LOGURU_LEVEL"),
)


def trtext2docx(
    source_text: Union[None, List[str], str] = None,
    trtext: Union[None, List[str], str] = None,
    alt_text: Union[None, List[str], str] = None,
    outfile: str = "",
    openfile: bool = True,
):
    """
    Generate docx for source_text, trtext, alt_text.

    Args:
    ----
    source_text: source text
    trtext: translated text
    alt_text: alternative translated text
    outfile: file path if True,
    openfile: try to open the saved file if True, default True

    Returns:
    -------
    docx.Document and saved file if file path supplied

    """
    if not source_text:
        source_text = []

    if not trtext:
        trtext = []

    if not alt_text:
        alt_text = []

    if isinstance(source_text, str):
        source_text = source_text.splitlines()

    if isinstance(trtext, str):
        trtext = trtext.splitlines()

    if isinstance(alt_text, str):
        alt_text = alt_text.splitlines()

    templ_dual = Path(__file__).parent / "templ_dual.docx"
    if templ_dual.exists():
        document = Document(templ_dual.as_posix())
        logger.info(f"Using {templ_dual=}")
    else:
        document = Document()

    # Normal style: built-in
    # document.styles["Normal"].font.name = "宋体"  # type: ignore
    # document.styles["Normal"].font.highlight_color = WD_COLOR_INDEX.YELLOW  # type: ignore
    document.styles["Normal"].font.size = Pt(12)
    document.styles["Normal"].paragraph_format.line_spacing = Pt(0)  # type: ignore

    # SrctextStyle
    srctext_style = document.styles.add_style("SrctextStyle", WD_STYLE_TYPE.PARAGRAPH)
    srctext_style.font.highlight_color = WD_COLOR_INDEX.YELLOW  # type: ignore
    srctext_style.font.size = Pt(10)
    document.styles["SrctextStyle"].paragraph_format.line_spacing = Pt(0)  # type: ignore

    # TrtextStyle
    trtext_style = document.styles.add_style("TrtextStyle", WD_STYLE_TYPE.PARAGRAPH)
    trtext_style.font.size = Pt(12)
    document.styles["TrtextStyle"].paragraph_format.line_spacing = Pt(12 + 4)  # type: ignore

    # Create Alttext style
    alttext_style = document.styles.add_style("AlttextStyle", WD_STYLE_TYPE.PARAGRAPH)
    alttext_font = alttext_style.font  # type: ignore
    alttext_font.size = Pt(10)
    document.styles["AlttextStyle"].paragraph_format.line_spacing = Pt(10 + 3)  # type: ignore

    alttext_font.highlight_color = WD_COLOR_INDEX.WHITE  # YELLOW WHITE GRAY_25
    alttext_font.color.rgb = RGBColor(0xFF, 0x0, 0xE0)

    paragraph = document.add_paragraph(style="Normal")
    for col0, col1, col2 in zip_longest(source_text, trtext, alt_text):
        if col0:
            # paragraph = document.add_paragraph(col0, style="Normal")
            paragraph = document.add_paragraph(col0, style="SrctextStyle")
        if col1:
            paragraph.paragraph_format.space_after = Pt(12)

            # paragraph = document.add_paragraph(col1, style="TrtextStyle")
            # paragraph = document.add_paragraph(style="TrtextStyle")
            paragraph = document.add_paragraph(style="Normal")
            run = paragraph.add_run(col1)
            # run.font.name = '楷体'
            run.font.name = "SimSun"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "NSimSun")
        if col2:
            paragraph.paragraph_format.space_after = Pt(4)
            # paragraph = document.add_paragraph(col2, style="AlttextStyle")
            paragraph = document.add_paragraph(style="AlttextStyle")
            run = paragraph.add_run(col2)
            run.font.name = "SimSun"
            # run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "NSimSun")
    if outfile:
        try:
            document.save(outfile)
            logger.info(f"Saved to {outfile=}")
        except Exception as exc:
            logger.warning(f" Cant save to {outfile=}: {exc}")
        if openfile:
            try:
                os.startfile(Path(outfile))
            except Exception as exc:
                logger.info(f"Cant open {outfile=}: {exc}")

    return document


async def main():  # pylint: disable=missing-function-docstring
    filepath = r"tests\test.txt"
    texts = loadtext(filepath)[:10]  # for testing

    y(filepath)
    console.print(f"{filepath=}")

    if isinstance(texts, str):
        texts = [texts]

    n_paras = len(texts)
    y(n_paras)
    console.print(f"{n_paras=} paragraphs")

    _ = "\n"
    console.print(f"""total: {len(texts)} paras, {_.join(texts[:3])}...""")

    n_workers = max(30, n_paras)

    then = monotonic()
    # dxtext = asyncio.run(batch_deeplx_tr(texts, n_workers=n_workers))
    dxtext = await batch_deeplx_tr(texts, n_workers=n_workers)
    # convert two-tuples to dict, in case there are missing entries
    dxtex_dict = dict(dxtext)
    col1 = [dxtex_dict.get(elm, "") for elm in range(n_paras)]
    console.print(f"done deeplx in {duration_human(monotonic() - then)}")

    # ofile = f"temp-{token_hex(3)}.docx"

    fname = Path(filepath).stem + f"-x{token_hex(3)}-tr.docx"

    ofile = (Path(filepath).parent / fname).absolute().as_posix()

    # trtext2docx(texts, col1, col1, ofile)
    trtext2docx(texts, col1, outfile=ofile)


if __name__ == "__main__":
    asyncio.run(main())
